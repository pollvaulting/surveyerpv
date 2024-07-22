import datetime
import io
import json
import os
import random
import re
import string
import threading
import time

from datetime import datetime
from datetime import date, timedelta

import jwt
import openai
import psycopg2
import pyotp
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from flask import Flask, jsonify, request, send_file
from flask_bcrypt import Bcrypt
from flask_cors import CORS
from flask_mail import Mail, Message

app = Flask(__name__)
cors = CORS(app)
bcrypt = Bcrypt()

openai.api_key = os.environ["OPENAI_API_KEY"]

APISECRET = os.environ.get('APISECRET')
hostsecret = os.environ['hostsecret']
portsecret = os.environ['portsecret']
usersecret = os.environ['usersecret']
passwordsecret = os.environ['passwordsecret']
databasesecret = os.environ['databasesecret']
mailsecret = os.environ['mailsecret']
mailpasswordsecret = os.environ['mailpasswordsecret']

db_params = {
    'host': hostsecret,
    'port': portsecret,
    'user': usersecret,
    'password': passwordsecret,
    'database': databasesecret
}

mail_params = {
    'email': mailsecret,  # Your email address for sending OTP
    'password': mailpasswordsecret,  # Your email password
    'server': 'smtp.gmail.com',  # Change based on your email provider
    'port': 587,
}
otp_secret_key = os.environ.get('APISECRET')
otp_expiry_seconds = 300  # OTP validity period in seconds

app.config['MAIL_SERVER'] = mail_params['server']
app.config['MAIL_PORT'] = mail_params['port']
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = mail_params['email']
app.config['MAIL_PASSWORD'] = mail_params['password']

users = []
profiles = []
polls = []
surveys = []


def get_db_connection():
  try:
    connection = psycopg2.connect(**db_params)
    return connection
  except Exception as e:
    print(f"Error: Unable to connect to the database. {str(e)}")
    return None


mail = Mail(app)


def generate_otp():

  totp = pyotp.TOTP(otp_secret_key, interval=otp_expiry_seconds)
  return totp.now()


def send_otp(email, otp):
  msg = Message('Your OTP for PollVault',
                sender='your-email@gmail.com',
                recipients=[email])
  msg.body = f'Your OTP is: {otp}'
  try:
    mail.send(msg)
    print(f"OTP sent to {email}")
    return True
  except Exception as e:
    print(f"Error sending OTP: {str(e)}")
    return False


def store_otp_in_database(email, otp):

  connection = psycopg2.connect(**db_params)
  cursor = connection.cursor()

  try:
    cursor.execute("INSERT INTO otps (email, otp) VALUES (%s, %s)",
                   (email, otp))
    connection.commit()
  except Exception as e:
    print(f"Error storing OTP in the database: {str(e)}")
    connection.rollback()
  finally:
    cursor.close()
    connection.close()


def verify_otp_in_database(email, otp):
  # Verify the OTP in the database
  connection = psycopg2.connect(**db_params)
  cursor = connection.cursor()

  try:
    cursor.execute("SELECT * FROM otps WHERE email = %s AND otp = %s",
                   (email, otp))
    result = cursor.fetchone()

    if result:
      return True
    else:
      return False
  except Exception as e:
    print(f"Error verifying OTP in the database: {str(e)}")
    return False
  finally:
    cursor.close()
    connection.close()


@app.route("/", methods=["GET"])
def index():
  return "API Online"


@app.route('/signup', methods=['POST'])
def sign_up():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  email = data.get('email')
  password = data.get('password')

  if not email or not password:
    return jsonify({
        'status': 'error',
        'error_content': 'Email and password are required'
    }), 400

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      cursor.execute("SELECT * FROM users WHERE email = %s", (email, ))
      existing_user = cursor.fetchone()
      if existing_user:
        return jsonify({
            'status': 'error',
            'error_content': 'Email already exists in the database'
        }), 400

      hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

      cursor.execute("INSERT INTO users (email, password) VALUES (%s, %s)",
                     (email, hashed_password))
      connection.commit()

      return jsonify({'status': 'success'}), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to register user'
      }), 500
    finally:
      cursor.close()
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/resetpassword', methods=['POST'])
def reset_password():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  email = data.get('email')
  password = data.get('password')

  if not email or not password:
    return jsonify({
        'status': 'error',
        'error_content': 'Email and password are required'
    }), 400

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      cursor.execute("SELECT * FROM users WHERE email = %s", (email, ))
      user = cursor.fetchone()
      if not user:
        return jsonify({
            'status': 'error',
            'error_content': 'Email doesn\'t exist in the database'
        }), 400

      if bcrypt.check_password_hash(user[2], password):
        return jsonify({
            'status': 'error',
            'error_content': 'This is the current password'
        }), 400

      hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
      cursor.execute("UPDATE users SET password = %s WHERE email = %s",
                     (hashed_password, email))
      connection.commit()

      return jsonify({'status': 'success'}), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to reset password'
      }), 500
    finally:
      cursor.close()
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


def validate_access_token(access_token):
  try:
    # Decode the access token
    decoded_token = jwt.decode(access_token,
                               'Pollvaultsecret',
                               algorithms=['HS256'])

    # Check if the token is not expired
    if datetime.utcnow() < datetime.utcfromtimestamp(decoded_token['exp']):
      return {
          'user_id': decoded_token['user_id'],
          'email': decoded_token['email']
      }
  except jwt.ExpiredSignatureError:
    pass  # Token has expired
  except jwt.InvalidTokenError:
    pass  # Invalid token or signature

  return None


def update_question_full(outline_id, connection):
  print("question full function")
  try:
    # Create a cursor
    cursor = connection.cursor()

    # Fetch rows based on outline_id
    query = "SELECT id, question_text, question_type, options FROM outline WHERE outline_id = %s"
    cursor.execute(query, (outline_id, ))
    rows = cursor.fetchall()

    for row in rows:
      id, question_text, question_type, options = row

      print(f"Processing row: {row}")

      if question_type == 'MCQ' and options:
        # If question_type is 'MCQ' and options exist, create question_full accordingly
        formatted_options = "\n".join(f"{chr(97 + i)}. {option}"
                                      for i, option in enumerate(options))
        #question_full = f"{question_text}\n {formatted_options}"
        print(f"formatted options for MCQ:{formatted_options}")
        update_query = "UPDATE outline SET formatted_options = %s WHERE id = %s"
        cursor.execute(update_query, (formatted_options, id))
      #else:

      # If question_type is not 'MCQ' or options don't exist, use question_text only
      #question_full = question_text
      #print(f"Question_full for non-MCQ: {question_full}")

      # Update the row with the new question_full using the unique id

    # Commit the changes
    connection.commit()

  finally:
    # Close the cursor
    cursor.close()


@app.route('/signin', methods=['POST'])
def sign_in():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  email = data.get('email')
  password = data.get('password')

  if not email or not password:
    return jsonify({
        'status': 'error',
        'error_content': 'Email and password are required'
    }), 400

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      cursor.execute("SELECT * FROM users WHERE email = %s", (email, ))
      user = cursor.fetchone()
      if not user:
        return jsonify({
            'status': 'error1',
            'error_content': 'Email doesn\'t exist in the database'
        }), 401

      if not bcrypt.check_password_hash(user[2], password):
        return jsonify({
            'status': 'error2',
            'error_content': 'Wrong password'
        }), 401

      # Generate JWT token
      token_payload = {
          'user_id': user[0],
          'email': user[1],
          'exp': datetime.utcnow() +
          timedelta(days=1)  # Token expiration (adjust as needed)
      }

      try:
        jwt_token = jwt.encode(token_payload,
                               'Pollvaultsecret',
                               algorithm='HS256')
        print(f"JWT Token: {jwt_token}")
        jwtd = jwt_token.decode('utf-8')
        # Insert data into the 'signedin' table
        cursor.execute(
            "INSERT INTO signedin (email, jwt_token) VALUES (%s, %s)",
            (email, jwtd))
        connection.commit()
      except Exception as e:
        print(f"Error encoding JWT: {str(e)}")

      # Return the token in the response
      return jsonify({
          'status': 'success',
          'token': jwt_token.decode(
              'utf-8')  # Convert bytes to string for JSON serialization
      }), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to sign in'
      }), 500
    finally:
      cursor.close()
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/profilesetup', methods=['POST'])
def profile_setup():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  #jwt_token = request.headers.get('JWTToken')

  #if not jwt_token:
  #  return jsonify({
  #      'status': 'error',
  #      'error_content': 'JWT token is required'
  #  }), 401

  # Validate the access token (implement your token validation logic)
  #user = validate_access_token(jwt_token)

  #if not user:
  #  return jsonify({
  #      'status': 'error',
  #      'error_content': 'Invalid JWT token'
  #  }), 401

  data = request.get_json()
  #email = data.get('email')
  salutation = data.get('salutation')
  first_name = data.get('firstname')
  last_name = data.get('lastname')
  industry = data.get('industry')
  location = data.get('location')
  phone = data.get('phone')

  #if not email:
  #  return jsonify({
  #      'status': 'error',
  #      'error_content': 'Email is required'
  #  }), 400

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      #cursor.execute("SELECT * FROM users WHERE email = %s", (email, ))
      #existing_user = cursor.fetchone()
      #if not existing_user:
      #  return jsonify({
      #      'status': 'error',
      #      'error_content': 'User not found'
      #  }), 404

      cursor.execute("SELECT * FROM profiles WHERE phone = %s", (phone, ))
      existing_phone = cursor.fetchone()
      if existing_phone:
        return jsonify({
            'status': 'error',
            'error_content': 'Phone number already in the database'
        }), 400

      cursor.execute(
          """
                INSERT INTO profiles (salutation, firstname, lastname, industry, location, phone)
                VALUES (%s, %s, %s, %s, %s, %s)

            """,
          (salutation, first_name, last_name, industry, location, phone))
      connection.commit()

      return jsonify({'status': 'success'}), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to set up user profile'
      }), 500
    finally:
      cursor.close()
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


# Dictionary to store outline IDs and their creation times
outline_creation_times = {}


def generate_outline_id(cursor):
  while True:
    outline_id = ''.join(random.choices(string.digits, k=6))

    cursor.execute("SELECT * FROM polls WHERE outline_id = %s", (outline_id, ))
    existing_poll = cursor.fetchone()

    if not existing_poll:
      return outline_id


# Create a function to generate questions and a random code using GPT-3.5 Turbo
def generate_questions(prompt):
  response = openai.ChatCompletion.create(
      model="gpt-4-1106-preview",
      messages=[
          {
              "role":
              "system",
              "content":
              "You are a helpful assistant that generates questions,type(MCQ or Free Text) from the given prompt and replies it in a json format and even if the content includes '\n' dont include it in the json response since we dont need it. It must reply in this format eg. : 'numberofquestions': len(questions)(2 in this case),'introduction': A brief introduction of the survey if it is given, otherwise generate it from the context, 'instruction': If specific instructions are mentioned, otherwise just answer the questions properly, 'time': appropriate time for answering the entire survey,'questions': (( 'question1': ('question': 'What is your favorite color?','type': 'MCQ','options': ['Red', 'Blue', 'Green']),('question2': ('question': 'What is your favorite programming language?','type': 'free text')) So understand this and reply it in a json format accordingly",
          },
          {
              "role": "user",
              "content": prompt
          },
      ],
  )
  generated_questions = response['choices'][0]['message']['content']
  return generated_questions


# Function to create the outline with actual generated questions
def create_outline(outline_id, connection, title, document):
  try:
    with connection.cursor() as cursor:
      # Add debug print statements to check the content of the 'document'
      print("Document:", document)
      generated_questions = generate_questions(document)
      print("Raw Response:", generated_questions)

      # Use regular expression to extract text between backticks and remove "json"
      match = re.search(r'```json([^`]+)```', generated_questions)

      if match:
        required_format = match.group(1)
        print(required_format)
      else:
        print("No match found")

      # Parse the generated_questions string as JSON
      generated_questions_response = json.loads(required_format)
      print("Parsed Response:", generated_questions_response)
      questions_data = generated_questions_response.get("questions", [])

      print("Questions Data:", questions_data)

      # Extract overall details
      number_of_questions = generated_questions_response.get(
          "numberofquestions", 0)
      survey_time = generated_questions_response.get("time", "unknown")

      print("survey time :" + survey_time)

      introduction = generated_questions_response.get("introduction",
                                                      "unknown")

      print("introduction :" + introduction)

      instruction = generated_questions_response.get("instruction", "unknown")

      print("instruction :" + instruction)

      for index, question_data in enumerate(questions_data, start=1):
        question_key = f"question{index}"
        print(f"Question {index}")

        question_text = question_data[question_key]['question']
        print(f"Question Text: {question_text}")

        question_type = question_data[question_key]['type']
        print(f"Question Type: {question_type}")

        options = json.dumps(question_data[question_key].get('options', []))
        print(f"Options: {options}")
        print()

        cursor.execute(
            """
                  INSERT INTO outline (outline_id, title, question_number, question_text, question_type, branching, options, importance, required, instruction, dynamic_followup, objective, max_no_of_questions, keywords_to_probe, things_to_avoid, example_questions, status)
                  VALUES (%s, %s, %s, %s, %s, FALSE, %s, 'normal', TRUE, NULL, FALSE, NULL, %s, NULL, NULL, NULL, 'Draft Ready')
                  """, (outline_id, title, index, question_text, question_type,
                        options, number_of_questions))

      connection.commit()

  except Exception as e:
    print(f"Error creating outline: {str(e)}")
  finally:
    pass


# Function to update poll status
def update_poll_status(outline_id, connection, title, document):
  # Simulate a delay of 1 minute before changing the status to 'created'
  time.sleep(5)

  cursor = connection.cursor()
  try:
    create_outline(outline_id, connection, title, document)
    update_question_full(outline_id, connection)

    cursor.execute(
        "UPDATE polls SET status = 'Draft Ready' WHERE outline_id = %s",
        (outline_id, ))
    connection.commit()

  except Exception as e:
    print(f"Error updating poll status: {str(e)}")
  finally:
    cursor.close()
    if connection:
      connection.close()


# Your existing upload_poll route
@app.route('/checkjwt', methods=['GET'])
def check_jwt():
  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  user = validate_access_token(jwt_token)

  if not user:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid JWT token'
    }), 401

  return jsonify({'status': 'success'}), 200


def skippernotes(document):
  # Instantiate your crew with a sequential process

  client = Anthropic()
  message = client.messages.create(
      model='claude-3-haiku-20240307',
      max_tokens=1024,
      messages=[
          {
              "role":
              "user",
              "content":
              f"Analyze the following uploaded document://starts here {document}  ends here//. The document includes an entire survey and there are some question skipping logic in some questions. Your task is to identify those questions which have skipper logic and to which question it is skipping to. Provide proper notes including the entire question and its skipping logic. Reply each skipping logic in a seperate line"
          },
      ]).content[0].text

  print(message)

  return (message)


def divideskippernotes(document):
  # Instantiate your crew with a sequential process

  client = Anthropic()
  message = client.messages.create(
      model='claude-3-haiku-20240307',
      max_tokens=1024,
      messages=[
          {
              "role":
              "user",
              "content":
              f""""Analyze the following skipper notes(skipper notes are basically a note on how  the selection of a particular option affects which the next question is going to be)://starts here {document}  ends here//. The document includes all the skipper logic of all the questions. Your task is to seperate them out as questions and provide them as a list to be easily extracted by regular expressions. Reply like this: 
       Q1. If option b, skip to question 15
       Q17. If option c, skip to question 20
       Q24. If option d, skip to question 29
So basically in a proper order which makes it easily extractable into different rows and columns
       """
          },
      ]).content[0].text

  print(message)

  return (message)


# Your existing upload_poll route
@app.route('/uploadpoll', methods=['POST'])
def upload_poll():
  api_secret = request.headers.get('APISECRET')
  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  user = validate_access_token(jwt_token)

  if not user:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid JWT token'
    }), 401

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  workspace = data.get('workspace')
  title = data.get('title')
  goal = data.get('goal')
  document = data.get('document')
  skipper = skippernotes(document)

  divideskipper = divideskippernotes(skipper)

  end_after_date = data.get('endafterdate')
  end_after_responses = data.get('endafterresponses')
  #email = data.get('email')
  geography = data.get('geography')
  education = data.get('education')
  industry = data.get('industry')
  visibility = data.get('visibility')

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      outline_id = generate_outline_id(cursor)

      cursor.execute(
          """
                INSERT INTO polls (outline_id, title, goal, document, endafterdate, endafterresponses,  geography, education, industry, visibility, status, jwt_token, workspace)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s, %s)
                """, (outline_id, title, goal, document, end_after_date,
                      end_after_responses, geography, education, industry,
                      visibility, 'Processing', jwt_token, workspace))
      cursor.execute(
          """
              INSERT INTO skipper_table (outline_id, skipperlogic)
              VALUES (%s, %s)
              """, (outline_id, skipper))

      connection.commit()

      # Define a regular expression pattern to match the question number and skipper logic
      pattern = r'Q(\d+)\. (.+?)\n'

      # Find all matches in the divideskipper string
      matches = re.findall(pattern, divideskipper)

      # Iterate over the matches and insert them into the PostgreSQL table
      for match in matches:
        cursor = connection.cursor()
        question_number = int(match[0])  # Convert question number to integer
        skipperlogic = match[1].strip()  # Remove leading/trailing spaces

        print("question: ", question_number)
        print("skipper logic: ", skipperlogic)
        # Assuming you have a specific outline ID
        # Insert into PostgreSQL table skipper_table_questions
        cursor.execute(
            """
              INSERT INTO skipper_table_questions (outline_id, question_number, skipperlogic)
              VALUES (%s, %s, %s)
              """, (outline_id, question_number, skipperlogic))
        connection.commit()

      # Store the creation time of the outline ID for later status update
      outline_creation_times[outline_id] = time.time()

      # Start a separate thread to update the status after 2 minutes
      threading.Thread(target=update_poll_status,
                       args=(outline_id, connection, title, document)).start()

      # Return the immediate response
      return jsonify({
          'status': 'success',
          'outline_id': outline_id,
          'processing_status': 'Processing'
      }), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to create poll'
      }), 500
    finally:
      if cursor:
        cursor.close()
  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


def generate_questions_zap(title, iwantto, tellusmore):
  prompt = f"{title}\n\nI want to {iwantto}\n\nTell us more: {tellusmore}"
  response = openai.ChatCompletion.create(
      model="gpt-4-1106-preview",
      messages=[
          {
              "role":
              "system",
              "content":
              """
              You are a helpful assistant that generates questions,type(MCQ or Free Text).// add the prompt here. Firstly , detect the 'industry' of this survey in 2 to 3 words (like a topic/field). Also detect the target responses and end date of this survey if it is mentioned, if not then reply in those fields as not mentioned. Same goes for whether they want the survey to be seen by anyone(public) or to select individuals/firms (private, reply as public or private only in this field).  Keep the number of questions between 5 to 15.  It must reply in this format eg. : 'numberofquestions': len(questions)(2 in this case),'instruction':if provided by the user otherwise just answer properly,'introduction': a brief introduction of the survey,'time': appropriate time for answering the entire survey,'industry': simple survey, 'Target Responses':100, 'Enddate': February 29 2024,visibility :'private', 'questions': (( 'question1': ('question': 'What is your favorite color?','type': 'MCQ','options': ['Red', 'Blue', 'Green']),('question2': ('question': 'What is your favorite programming language?','type': 'free text')). Keep the numberofquestions below 10. Also return it in a proper json format so that i can dump it directly in json.dump An example result is this: {'numberofquestions': 17, 'time': '30 minutes', 'questions': [{'question1': {'question': 'Which of the following best describes you?', 'type': 'MCQ', 'options': ['About to graduate', 'Recent graduate (within last 2 years)']}}, {'question2': {'question': 'Do you feel that your academic program has adequately equipped you to apply leadership skills in a real-world setting?', 'type': 'MCQ', 'options': ['Strongly agree', 'Agree', 'Neutral', 'Disagree', 'Strongly disagree']}}, {'question3': {'question': 'What could your graduate education have done differently to better prepare you for leadership roles?', 'type': 'free text'}}, {'question4': {'question': 'Can you select the types of leadership development opportunities you have received during your graduate education?', 'type': 'MCQ', 'options': ['Leadership courses', 'Private courses (e.g., speaking) sponsored by the university', 'Mentorship', 'Leadership competitions', 'Self-study', 'Other']}}, {'question5': {'question': 'Which were most useful and why? Which were least useful and why?', 'type': 'free text'}}, {'question6': {'question': 'After graduation, which healthcare institution do you aim to work for?', 'type': 'MCQ', 'options': ['Physician group', 'Health system', 'Health plan', 'Life sciences', 'Local or state agency', 'Federal agency', 'Non-profit', 'Other']}}, {'question7': {'question': 'Please briefly describe any previous leadership roles/experience.', 'type': 'free text'}}, {'question8': {'question': 'Based on feedback from internships or academic projects, what do you believe are your leadership strengths?', 'type': 'free text'}}, {'question9': {'question': 'Which leadership areas have been identified for improvement?', 'type': 'free text'}}, {'question10': {'question': 'In your opinion, what are the three most important leadership qualities?', 'type': 'free text'}}, {'question11': {'question': 'Which institution and which degree are you graduating from or have graduated from?', 'type': 'free text'}}, {'question12': {'question': 'Would you be willing to connect for a follow-up interview?', 'type': 'free text'}}, {'question13': {'question': 'Any other thoughts on leadership? Graduate education?', 'type': 'free text'}}, {'question14': {'question': 'What kind of organization do you work for? (For example, health system, government)', 'type': 'free text'}}, {'question15': {'question': 'What is your current role?', 'type': 'free text'}}, {'question16': {'question': 'Do you feel you have adequate leadership development opportunities on the job? Please explain.', 'type': 'free text'}}, {'question17': {'question': 'How closely do you believe your academic training aligns with the real-world expectations of leadership based on your experience to-date?', 'type': 'MCQ', 'options': ['Completely misaligned', 'Somewhat misaligned', 'Neutral', 'Somewhat aligned', 'Very aligned']}}], 'suggested_question': {'question18': {'question': 'SUGGESTED: How has the evolving healthcare landscape influenced your views on leadership in healthcare management?', 'type': 'free text'}}}
""",
          },
          {
              "role": "user",
              "content": prompt
          },
      ],
  )
  generated_questions = response['choices'][0]['message']['content']
  return generated_questions


# Function to create the outline with actual generated questions
def create_zapoutline(outline_id, connection, title, iwantto, tellusmore):
  try:
    with connection.cursor() as cursor:
      # Add debug print statements to check the content of the 'document'

      generated_questions = generate_questions_zap(title, iwantto, tellusmore)
      print("Raw Response:", generated_questions)

      # Use regular expression to extract text between backticks and remove "json"
      match = re.search(r'```json([^`]+)```', generated_questions)

      if match:
        required_format = match.group(1)
        print(required_format)
      else:
        print("No match found")

      # Parse the generated_questions string as JSON
      generated_questions_response = json.loads(required_format)

      print("Parsed Response:", generated_questions_response)
      questions_data = generated_questions_response.get("questions", [])

      print("Questions Data:", questions_data)

      # Extract overall details
      number_of_questions = generated_questions_response.get(
          "numberofquestions", 0)
      survey_time = generated_questions_response.get("time", "unknown")
      survey_industry = generated_questions_response.get("industry", "unknown")
      targetresponses = generated_questions_response.get("Target Responses", 0)
      targetresponses = 100
      visibility = generated_questions_response.get("visibility", "unknown")

      introduction = generated_questions_response.get("introduction",
                                                      "unknown")

      print("introduction :" + introduction)

      instruction = generated_questions_response.get("instruction", "unknown")

      print("instruction :" + instruction)

      Enddate = generated_questions_response.get("End date", date(2001, 1, 1))
      Enddate = date(2024, 2, 29)
      cursor.execute(
          """
          UPDATE polls
          SET goal = %s, document = %s, endafterdate = %s, endafterresponses = %s,
              geography = %s, education = %s, industry = %s, visibility = %s
          WHERE outline_id = %s
          """, (survey_industry, "zap", Enddate, targetresponses, "ZAP", "ZAP",
                survey_industry, visibility, outline_id))

      for index, question_data in enumerate(questions_data, start=1):
        question_key = f"question{index}"
        print(f"Question {index}")

        question_text = question_data[question_key]['question']
        print(f"Question Text: {question_text}")

        question_type = question_data[question_key]['type']
        print(f"Question Type: {question_type}")

        options = json.dumps(question_data[question_key].get('options', []))
        print(f"Options: {options}")
        print()

        cursor.execute(
            """
                  INSERT INTO outline (outline_id, title, question_number, question_text, question_type, branching, options, importance, required, instruction, dynamic_followup, objective, max_no_of_questions, keywords_to_probe, things_to_avoid, example_questions, status)
                  VALUES (%s, %s, %s, %s, %s, FALSE, %s, 'normal', TRUE, NULL, FALSE, NULL, %s, NULL, NULL, NULL, 'Draft Ready')
                  """, (outline_id, title, index, question_text, question_type,
                        options, number_of_questions))

      connection.commit()

  except Exception as e:
    print(f"Error creating outline: {str(e)}")
  finally:
    pass


# Function to update poll status
def update_zappoll_status(outline_id, connection, title, iwantto, tellusmore):
  # Simulate a delay of 1 minute before changing the status to 'created'
  time.sleep(5)

  cursor = connection.cursor()
  try:
    create_zapoutline(outline_id, connection, title, iwantto, tellusmore)
    update_question_full(outline_id, connection)

    cursor.execute(
        "UPDATE polls SET status = 'Draft Ready' WHERE outline_id = %s",
        (outline_id, ))
    connection.commit()

  except Exception as e:
    print(f"Error updating poll status: {str(e)}")
  finally:
    cursor.close()
    if connection:
      connection.close()


@app.route('/uploadzappoll', methods=['POST'])
def upload_zappoll():
  api_secret = request.headers.get('APISECRET')
  data = request.get_json()
  workspace = data.get('workspace')
  title = data.get('title')
  iwantto = data.get('iwantto')
  tellusmore = data.get('document')
  #questions_reponse = generate_questions_zap(title, iwantto, tellusmore)

  #return questions_reponse

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()
      outline_id = generate_outline_id(cursor)

      cursor.execute(
          """
                INSERT INTO polls (outline_id, title, jwt_token, status,workspace)
                VALUES (%s, %s, %s,%s, %s)
                """, (outline_id, title, jwt_token, 'Processing', workspace))

      connection.commit()

      # Store the creation time of the outline ID for later status update
      outline_creation_times[outline_id] = time.time()

      # Start a separate thread to update the status after 2 minutes
      threading.Thread(target=update_zappoll_status,
                       args=(outline_id, connection, title, iwantto,
                             tellusmore)).start()

      # Return the immediate response
      return jsonify({
          'status': 'success',
          'outline_id': outline_id,
          'processing_status': 'Processing'
      }), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to create poll'
      }), 500
    finally:
      if cursor:
        cursor.close()
  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


@app.route('/getoutline', methods=['POST'])
def get_outline():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  #user = validate_access_token(jwt_token)

  #if not user:
  #  return jsonify({
  #      'status': 'error',
  #      'error_content': 'Invalid JWT token'
  #  }), 401

  data = request.get_json()
  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Fetch data from the outline table based on outline_id
        cursor.execute(
            """
                  SELECT * FROM outline
                  WHERE outline_id = %s
                  ORDER BY question_number
                  """, (outline_id, ))
        outline_data = cursor.fetchall()

        if outline_data:
          # Get the column names from cursor.description
          column_names = [desc[0] for desc in cursor.description]

          cursor.execute(
              """
                      SELECT goal, endafterdate, endafterresponses, geography, education, industry, visibility
                      FROM polls
                      WHERE outline_id = %s
                      """, (outline_id, ))
          survey_data = cursor.fetchall()

          # Construct the response JSON with the fetched outline data
          response_json = {
              'outline': {
                  'numberofquestions':
                  len(outline_data),
                  'questions': [],
                  'time':
                  '10 mins',
                  'title':
                  outline_data[0][column_names.index('title')],
                  'status':
                  outline_data[0][column_names.index('status')],
                  'outline_id':
                  outline_data[0][column_names.index('outline_id')],
                  'id':
                  outline_data[0][column_names.index('id')],
                  'goal':
                  None,
                  'endafterdate':
                  None,
                  'geography':
                  None,
                  'education':
                  None,
                  'industry':
                  None,
                  'visibility':
                  None,
                  'introduction':
                  outline_data[0][column_names.index('instruction')],
                  'instruction':
                  outline_data[0][column_names.index('instruction')]
              },
              'outlinestatus': outline_data[0][column_names.index('status')],
              'version': '1.05',
              'status': 'success'
          }

          if survey_data and len(survey_data) == 1:
            # Get the column names from cursor.description for survey_data
            column_names_survey = [desc[0] for desc in cursor.description]

            input_date = survey_data[0][column_names_survey.index(
                'endafterdate')]

            if input_date:
              if isinstance(input_date, (date, datetime)):
                formatted_date = input_date.strftime("%d %b %Y")

            else:
              # If it's a string, then use strptime and strftime
              formatted_date = 'NULL'
              try:
                parsed_date = datetime.strptime(input_date,
                                                "%a, %d %b %Y %H:%M:%S %Z")
                formatted_date = parsed_date.strftime("%d %b %Y")

              except ValueError as e:
                print(f"Error parsing date {input_date}: {e}")

            # Update the response JSON with survey data
            response_json['outline'].update({
                'goal':
                survey_data[0][column_names_survey.index('goal')],
                'endafterdate':
                formatted_date,
                'geography':
                survey_data[0][column_names_survey.index('geography')],
                'education':
                survey_data[0][column_names_survey.index('education')],
                'industry':
                survey_data[0][column_names_survey.index('industry')],
                'visibility':
                survey_data[0][column_names_survey.index('visibility')]
            })

          # Sort outline_data by question_number in ascending order
          sorted_outline_data = sorted(
              outline_data,
              key=lambda row: int(row[column_names.index('question_number')]))

          # Iterate through fetched data and structure the questions
          for row in sorted_outline_data:
            if row[column_names.index('dynamic_followup')] == 'true':
              dfu = True
            else:
              dfu = False

            if row[column_names.index('required')] == 'true':
              dfr = True
            else:
              dfr = False

            question_number = int(row[column_names.index('question_number')])

            # Query 'skipper_table_questions' for skipper logic
            cursor.execute(
                """
                SELECT skipperlogic
                FROM skipper_table_questions
                WHERE outline_id = %s AND question_number = %s
                """, (outline_id, question_number))
            skipper_logic_row = cursor.fetchone()
            skipper_logic = skipper_logic_row[0] if skipper_logic_row else None

            question_data = {
                "question_number":
                row[column_names.index('question_number')],
                "branching":
                row[column_names.index('branching')],
                "question":
                row[column_names.index('question_text')],
                "options":
                row[column_names.index('formatted_options')],
                "type":
                row[column_names.index('question_type')],
                "importance":
                row[column_names.index('importance')],
                "required":
                dfr,
                "dynamic_followup":
                dfu,
                "objective":
                row[column_names.index('objective')],
                "max_no_of_questions":
                row[column_names.index('max_no_of_questions')],
                "keywords_to_probe":
                row[column_names.index('keywords_to_probe')],
                "things_to_avoid":
                row[column_names.index('things_to_avoid')],
                "example_questions":
                row[column_names.index('example_questions')],
                "allow_others":
                row[column_names.index('allow_others')],
                "max_no_of_choices":
                row[column_names.index('max_no_of_choices')],
                "skipper_logic":
                skipper_logic,
                # Include other columns as needed
            }

            response_json['outline']['questions'].append(question_data)

          return jsonify(response_json), 200
        else:
          return jsonify({
              'status': 'error',
              'error_content': 'Outline ID does not exist'
          }), 404

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to fetch outline data. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


def reverse_format_options(formatted_options):
  if formatted_options:
    # Split the formatted options into individual lines
    options_lines = formatted_options.split('\n')

    # Extract the options from each line
    options = [line.split('. ')[1] for line in options_lines]

    return options
  else:
    return None


@app.route('/saveoutline', methods=['POST'])
def save_outline():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  user = validate_access_token(jwt_token)
  if not user:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid JWT token'
    }), 401

  data = request.get_json()
  outline_data = data.get('outline')

  if not outline_data:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Extract outline information
        title = outline_data.get('title')
        outline_id = outline_data.get('outline_id')
        status = outline_data.get('status')

        # Delete existing rows for the given outline_id
        cursor.execute("DELETE FROM outline WHERE outline_id = %s",
                       (outline_id, ))

        cursor.execute(
            """
            Delete from skipper_table_questions where outline_id = %s
          """, (outline_id, ))

        cursor.execute(
            """
              UPDATE polls
              SET title = %s
              WHERE outline_id = %s
              """, (title, outline_id))

        # Insert new rows into the outline table
        for question_data in outline_data.get('questions', []):

          print(question_data['options'])

          actual_options = reverse_format_options(question_data['options'])
          cursor.execute(
              """
                        INSERT INTO outline (
                            title, question_number, importance, required, instruction,
                            dynamic_followup, objective, max_no_of_questions,
                            keywords_to_probe, things_to_avoid, example_questions,allow_others,max_no_of_choices,
                            question_text,formatted_options, question_type, branching,outline_id, status,options
                        )
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s)
                        """,
              (title, question_data.get('question_number'),
               question_data.get('importance', 'normal'),
               question_data.get('required',
                                 'TRUE'), question_data.get('instruction'),
               question_data.get('dynamic_followup',
                                 'FALSE'), question_data.get('objective'),
               question_data.get('max_no_of_questions'),
               question_data.get('keywords_to_probe'),
               question_data.get('things_to_avoid'),
               question_data.get('example_questions'),
               question_data.get('allow_others'),
               question_data.get('max_no_of_choices'),
               question_data['question'], question_data['options'],
               question_data['type'], question_data.get('branching', 'FALSE'),
               outline_id, status, json.dumps(actual_options)))

          qn = int(question_data.get('question_number'))
          print("hi")
          cursor.execute(
              """
              INSERT INTO skipper_table_questions (outline_id, question_number, skipperlogic)
              VALUES (%s, %s, %s)
              """,
              (outline_id, qn, question_data.get('skipper_logic', 'FALSE')))

        connection.commit()

        return jsonify({
            'status': 'success',
            'message': 'Outline saved successfully'
        }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to save outline. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/getpolls', methods=['POST'])
def get_polls():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Fetch data from the outline table based on outline_id
        cursor.execute(
            """
                  SELECT title, goal, endafterdate, endafterresponses, geography, education, industry, visibility,instruction,introduction
                      FROM polls
                      WHERE outline_id = %s
                      """, (outline_id, ))
        survey_data = cursor.fetchall()

        if survey_data:
          # Get the column names from cursor.description
          column_names = [desc[0] for desc in cursor.description]

          input_date = survey_data[0][column_names.index('endafterdate')]
          introduction = survey_data[0][column_names.index('introduction')]
          print(introduction)
          if input_date:
            if isinstance(input_date, (date, datetime)):
              formatted_date = input_date.strftime("%d %b %Y")

          else:
            # If it's a string, then use strptime and strftime
            formatted_date = 'NULL'
            try:
              parsed_date = datetime.strptime(input_date,
                                              "%a, %d %b %Y %H:%M:%S %Z")
              formatted_date = parsed_date.strftime("%d %b %Y")

            except ValueError as e:
              print(f"Error parsing date {input_date}: {e}")

          if introduction == 'null':
            introduction = ""

          # Update the response JSON with survey data

          # Construct the response JSON with the fetched outline data
          response_json = {
              'survey': {
                  'title':
                  survey_data[0][column_names.index('title')],
                  'version':
                  '1.05',
                  'goal':
                  survey_data[0][column_names.index('goal')],
                  'endafterdate':
                  formatted_date,
                  'endafterresponses':
                  survey_data[0][column_names.index('endafterresponses')],
                  'geography':
                  survey_data[0][column_names.index('geography')],
                  'education':
                  survey_data[0][column_names.index('education')],
                  'industry':
                  survey_data[0][column_names.index('industry')],
                  'visibility':
                  survey_data[0][column_names.index('visibility')],
                  'introduction':
                  "",
                  'instruction':
                  ""
              },
              'status': 'success'
          }

          return jsonify(response_json), 200
        else:
          return jsonify({
              'status': 'error',
              'error_content': 'Outline ID does not exist'
          }), 404

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to fetch outline data. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/savepolls', methods=['POST'])
def save_polls():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  user = validate_access_token(jwt_token)
  if not user:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid JWT token'
    }), 401

  data = request.get_json()
  outline_data = data.get('outline')

  if not outline_data:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Extract outline information

        outline_id = outline_data.get('outline_id')

        goal = outline_data.get('goal')
        endafterdate = outline_data.get('endafterdate')
        endafterresponses = outline_data.get('endafterresponses')
        geography = outline_data.get('geography')
        education = outline_data.get('education')
        industry = outline_data.get('industry')
        visibility = outline_data.get('visibility')
        introduction = outline_data.get('introduction')
        instruction = outline_data.get('instruction')

        cursor.execute("DELETE FROM polls WHERE outline_id = %s",
                       (outline_id, ))

        cursor.execute(
            """
              INSERT INTO polls (outline_id, goal, endafterdate, endafterresponses, geography, education, industry, visibility, introduction, instruction)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
              """,
            (outline_id, goal, endafterdate, endafterresponses, geography,
             education, industry, visibility, introduction, instruction))

        connection.commit()

        return jsonify({
            'status': 'success',
            'message': 'Outline saved successfully'
        }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to save outline. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


def greetise(question_text):
  response = openai.ChatCompletion.create(
      model="gpt-4o",
      messages=[
          {
              "role":
              "system",
              "content":
              """
              You are a helpful assistant that slightly changes the tone of the question to a more humanlike tone but it must be very minimal changes. Note that these questions may not be the first questions in the list, so dont put 'Hey there', or any sort of greetings like that. and make sure that the length of the question remains the same if its a example 25 words question then the new generated question must not be more than 35 words.
              Example: Which of the following best describes you - would turn into Just wanted to ask which one of the below desribes you best? and Which leadership areas have been identified for improvement? to Which areas of the leadership do u think you need to focus on?, etc
""",
          },
          {
              "role": "user",
              "content": question_text
          },
      ],
  )
  generated_question = response['choices'][0]['message']['content']
  return generated_question


# Function to update poll status
def update_poll_status_published(outline_id, connection):
  # Simulate a delay of 1 minute before changing the status to 'created'
  time.sleep(5)

  cursor = connection.cursor()
  try:

    cursor.execute(
        "UPDATE polls SET status = 'Published' WHERE outline_id = %s",
        (outline_id, ))
    connection.commit()

  except Exception as e:
    print(f"Error updating poll status: {str(e)}")
  finally:
    cursor.close()
    if connection:
      connection.close()


@app.route('/getworld', methods=['POST'])
def getworld():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Fetch data from the outline table based on outline_id
        cursor.execute(
            """
                  select survey_code from published where outline_id = %s limit 1
                  """, (outline_id, ))
        res = cursor.fetchall()
        column_names = [desc[0] for desc in cursor.description]

        survey_code = res[0][column_names.index('survey_code')]

      return jsonify({
          'status': 'success',
          'message': 'This is a published survey',
          'survey_code': survey_code
      }), 200

    except:
      return jsonify({
          'status': 'error',
          'error_content': 'outline id not published'
      }), 500


@app.route('/publishsurvey', methods=['POST'])
def publish_survey():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  outline_data = data.get('outline')

  if not outline_data:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Extract outline information
        title = outline_data.get('title')
        outline_id = outline_data.get('outline_id')
        status = outline_data.get('status')

        # Delete existing rows for the given outline_id
        cursor.execute("DELETE FROM outline WHERE outline_id = %s",
                       (outline_id, ))
        cursor.execute(
            """
            UPDATE polls
            SET title = %s
            WHERE outline_id = %s
            """, (title, outline_id))

        # Insert new rows into the outline table
        for question_data in outline_data.get('questions', []):
          actual_options = reverse_format_options(question_data.get('options'))
          question_text = greetise(question_data['question'])
          cursor.execute(
              """
                        INSERT INTO outline (
                            title, question_number, importance, required, instruction,
                            dynamic_followup, objective, max_no_of_questions,
                            keywords_to_probe, things_to_avoid, example_questions,allow_others,max_no_of_choices,
                            question_text,formatted_questions, formatted_options, question_type, branching,outline_id, status,options
                        )
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s)
                        """,
              (title, question_data.get('question_number'),
               question_data.get('importance', 'normal'),
               question_data.get('required',
                                 'TRUE'), question_data.get('instruction'),
               question_data.get('dynamic_followup',
                                 'FALSE'), question_data.get('objective'),
               question_data.get('max_no_of_questions'),
               question_data.get('keywords_to_probe'),
               question_data.get('things_to_avoid'),
               question_data.get('example_questions'),
               question_data.get('allow_others'),
               question_data.get('max_no_of_choices'),
               question_data['question'], question_text,
               question_data['options'], question_data['type'],
               question_data.get('branching', 'FALSE'), outline_id, status,
               json.dumps(actual_options)))

        # Generate a unique 6-digit survey code
        survey_code = generate_unique_survey_code(cursor)

        # Insert into the published table
        cursor.execute(
            """
                    INSERT INTO published (outline_id, survey_code)
                    VALUES (%s, %s)
                    """, (outline_id, survey_code))

        connection.commit()

        update_poll_status_published(outline_id, connection)

        return jsonify({
            'status': 'success',
            'message': 'Survey published successfully',
            'survey_code': survey_code
        }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to publish survey. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/publishsurveydirect', methods=['POST'])
def publish_survey_direct():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()

  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Outline id needed'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:

        # Generate a unique 6-digit survey code
        survey_code = generate_unique_survey_code(cursor)

        # Insert into the published table
        cursor.execute(
            """
                    INSERT INTO published (outline_id, survey_code)
                    VALUES (%s, %s)
                    """, (outline_id, survey_code))

        connection.commit()

        update_poll_status_published(outline_id, connection)

        return jsonify({
            'status': 'success',
            'message': 'Survey published successfully',
            'survey_code': survey_code
        }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to publish survey. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


def generate_unique_survey_code(cursor):
  # Generate a unique 6-digit survey code that doesn't exist in the published table
  while True:
    survey_code = ''.join(random.choices(string.digits, k=6))
    cursor.execute("SELECT COUNT(*) FROM published WHERE survey_code = %s",
                   (survey_code, ))
    count = cursor.fetchone()[0]
    if count == 0:
      return survey_code


@app.route('/responsequestions', methods=['POST'])
def response_questions():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')

  if not survey_code:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Check if survey code exists in the published table
        cursor.execute(
            "SELECT outline_id FROM published WHERE survey_code = %s",
            (survey_code, ))
        result = cursor.fetchone()

        if result:
          outline_id = result[0]

          # Fetch data from the outline table based on outline_id
          cursor.execute(
              """
                        SELECT * FROM outline
                        WHERE outline_id = %s
                        ORDER BY question_number
                        """, (outline_id, ))
          outline_data = cursor.fetchall()

          if outline_data:
            # Get the column names from cursor.description
            column_names = [desc[0] for desc in cursor.description]

            # Construct the response JSON with the fetched outline data
            response_json = {
                'outline': {
                    'numberofquestions': len(outline_data),
                    'questions': [],
                    'time': '10 minutes',
                    'title': outline_data[0][column_names.index('title')],
                    'status': outline_data[0][column_names.index('status')],
                    'outline_id':
                    outline_data[0][column_names.index('outline_id')],
                    'id': outline_data[0][column_names.index('id')],
                },
                'outlinestatus': outline_data[0][column_names.index('status')],
                'status': 'success'
            }

            # Iterate through fetched data and structure the questions
            for row in outline_data:
              question_data = {
                  "question_number":
                  row[column_names.index('question_number')],
                  "branching":
                  row[column_names.index('branching')],
                  "question":
                  row[column_names.index('question_text')],
                  "type":
                  row[column_names.index('question_type')],
                  "options":
                  row[column_names.index('options')],
                  "importance":
                  row[column_names.index('importance')],
                  "required":
                  row[column_names.index('required')],
                  "instruction":
                  row[column_names.index('instruction')],
                  "dynamic_followup":
                  row[column_names.index('dynamic_followup')],
                  "objective":
                  row[column_names.index('objective')],
                  "max_no_of_questions":
                  row[column_names.index('max_no_of_questions')],
                  "keywords_to_probe":
                  row[column_names.index('keywords_to_probe')],
                  "things_to_avoid":
                  row[column_names.index('things_to_avoid')],
                  "example_questions":
                  row[column_names.index('example_questions')],
                  # Include other columns as needed
              }

              response_json['outline']['questions'].append(question_data)

            return jsonify(response_json), 200
          else:
            return jsonify({
                'status': 'error',
                'error_content': 'Outline ID does not exist'
            }), 404

        else:
          return jsonify({
              'status': 'error',
              'error_content': 'Survey code not found'
          }), 404

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to fetch outline data. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/saveanswers', methods=['POST'])
def save_answers():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')
  question_data = data.get('question')
  answer_data = data.get('answer')

  if not survey_code or not question_data or not answer_data:
    return jsonify({
        'status': 'error',
        'error_content': 'Incomplete data provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Insert answer into the answers table
        cursor.execute(
            """
                    INSERT INTO answers (
                        survey_code, question_number, question, type, options, answer
                    )
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
            (survey_code, question_data.get('question_number'),
             question_data.get('question'), question_data.get('type'),
             json.dumps(question_data.get('options', [])), answer_data))

        connection.commit()

        return jsonify({
            'status': 'success',
            'message': 'Answers saved successfully'
        }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to save answers. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/sendotp', methods=['POST'])
def send_otp_route():
  email = request.json.get('email')

  otp = generate_otp()
  send_otp(email, otp)

  store_otp_in_database(email, otp)

  return jsonify({'status': 'success'}), 200


@app.route('/verifyotp', methods=['POST'])
def verify_otp_route():
  email = request.json.get('email')
  user_otp = request.json.get('otp')

  if verify_otp_in_database(email, user_otp):
    return jsonify({'status': 'success'}), 200
  else:
    return jsonify({'status': 'error', 'error_content': 'Wrong OTP'}), 401


# Get Outline Status
@app.route('/getoutlinestatus', methods=['POST'])
def get_outline_status():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Fetch outline status based on outline_id
        cursor.execute(
            """
                    SELECT status FROM outline
                    WHERE outline_id = %s
                    """, (outline_id, ))
        outline_status = cursor.fetchone()

        if outline_status:
          return jsonify({
              'status': 'success',
              'outlinestatus': outline_status[0]
          }), 200
        else:
          return jsonify({
              'status': 'error1',
              'error_content': 'Outline ID does not exist'
          }), 404

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to fetch outline status. Error: {str(e)}'
      }), 500

    finally:
      if connection:
        connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/getfirstquestion', methods=['POST'])
def get_first_question():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')

  responder_id = data.get('responder_id')
  if not responder_id:
    responder_id = '0'

  if not survey_code:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Check if survey code exists in the published table
        cursor.execute(
            "SELECT outline_id FROM published WHERE survey_code = %s",
            (survey_code, ))
        result = cursor.fetchone()

        if result:
          # Survey code found, now retrieve the first question details from the outline table
          outline_id = result[0]
          cursor.execute(
              """
                        SELECT question_number, formatted_questions, question_type, options,max_no_of_choices, required, question_text
                        FROM outline
                        WHERE outline_id = %s AND question_number = '1'
                    """, (outline_id, ))
          question_result = cursor.fetchone()

          if question_result:
            question_number, formatted_questions, question_type, options, max_no_of_choices, required, question_text = question_result
            if required == 'true':
              skipper = False
            else:
              skipper = True

            # Return the first question details
            if formatted_questions is not None:
              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': formatted_questions,
                  'type': question_type,
                  'options': options,
                  'max_no_of_choices': max_no_of_choices,
                  'skipquestion': skipper,
                  'total_time': '45',
                  'responder_id': responder_id,
                  'introduction':
                  'Introduction: We are seeking to understand the perceptions of healthcare management students and recent graduates regarding leadership development in healthcare.  Little more context ',
                  'instruction': 'Answer each questions carefully'
              }), 200
            else:
              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': question_text,
                  'type': question_type,
                  'options': options,
                  'max_no_of_choices': max_no_of_choices,
                  'skipquestion': skipper,
                  'total_time': '45',
                  'responder_id': responder_id,
                  'introduction':
                  'Introduction: We are seeking to understand the perceptions of healthcare management students and recent graduates regarding leadership development in healthcare.  Little more context ',
                  'instruction': 'Answer each questions carefully'
              }), 200

        # Survey code not found
        return jsonify({
            'status': 'error',
            'error_content': 'Survey code not found'
        }), 404

    except Exception as e:
      print(f"Error getting first question: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to retrieve the first question'
      }), 500

    finally:
      connection.close()

  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


@app.route('/getrespondent', methods=['POST'])
def getrespondent():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')

  responder_id = data.get('responder_id')
  if not responder_id:
    responder_id = '0'

  if not survey_code:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Check if survey code exists in the published table
        cursor.execute(
            "SELECT outline_id FROM published WHERE survey_code = %s",
            (survey_code, ))
        result = cursor.fetchone()

        if result:
          # Survey code found, now retrieve the first question details from the outline table
          outline_id = result[0]
          cursor.execute(
              """
                        SELECT question_number, question_text, question_type, options,max_no_of_choices, required
                        FROM outline
                        WHERE outline_id = %s AND question_number = '1'
                    """, (outline_id, ))
          question_result = cursor.fetchone()

          if question_result:
            question_number, question_text, question_type, options, max_no_of_choices, required = question_result
            if required == 'true':
              skipper = False
            else:
              skipper = True

          if outline_id == '556782':
            return jsonify({
                'introduction':
                'We are seeking to understand the entire hospital experience for your length of stay so we can make improvements to our system properly',
                'instruction': 'Answer each questions carefully',
                'totalquestion': '29',
                'title': 'Healthcare Leadership Survey',
                'totaltime': '45'
            }), 200
            # Return the first question details
          else:
            return jsonify({
                'introduction':
                ' We are seeking to understand the perceptions of healthcare management students and recent graduates regarding leadership development in healthcare.  Little more context ',
                'instruction': 'Answer each questions carefully',
                'totalquestion': '23',
                'title': 'HCAPS Survey',
                'totaltime': '45'
            }), 200

        # Survey code not found
        return jsonify({
            'status': 'error',
            'error_content': 'Survey code not found'
        }), 404

    except Exception as e:
      print(f"Error getting question: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to retrieve the intro question'
      }), 500

    finally:
      connection.close()

  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


def answerify(answer, options):
  response = openai.ChatCompletion.create(
      model="gpt-4o",
      messages=[
          {
              "role":
              "system",
              "content":
              f"""
              You are a helpful assistant that checks this transcript of a user's answer to a mcq question to one of these options:{options} and matches them to the correct answer and returns the exact answer included in the option. If it matches neither (double recheck), then return "No match found" but make sure that it must not be from the options before u return this, this is a last resort answer. And also dont include any sort of comments in the answer, just the exact option with the key, (a. , b., etc).
""",
          },
          {
              "role": "user",
              "content": answer
          },
      ],
  )
  generated_question = response['choices'][0]['message']['content']
  return generated_question


def dynamic_voice_question(formatted_questions, answer):
  response = openai.ChatCompletion.create(
      model="gpt-4o",
      messages=[
          {
              "role":
              "system",
              "content":
              f"""
            You are a helpful assistant which reformats this question :{formatted_questions} and adds tone to it and make it related as if it is continuous to the previous question's answer given below. Note: only return the question ,nothing else. 
  """,
          },
          {
              "role": "user",
              "content": answer
          },
      ],
  )
  generated_question = response['choices'][0]['message']['content']
  return generated_question


def whetherskip(question, answer, skipper_logic):
  # Instantiate your crew with a sequential process

  client = Anthropic()
  message = client.messages.create(
      model='claude-3-sonnet-20240229',
      #model = 'claude-3-haiku-20240307',
      max_tokens=2048,
      messages=[
          {
              "role":
              "user",
              "content":
              f""" Follow the instructions carefully and provide the output in the specified format.Analyze the following question and answer and the notes of skipper logic. Skipper logic means when there is a particular answer to a question, it may lead to skipping the next few questions. :\n\nQuestion: //starts here {question} ends here//Answer: //starts here {answer} ends here//Skipper Logic: //starts here {skipper_logic} ends here//Instructions:1. Check if the provided question and answer match any of the skipper logic conditions exactly.  \n2. If a skip condition is met, respond with only the question number to skip to (e.g., '5'). Do not include any explanations or additional text.\n3. If no skip condition is met, respond with only the word 'nope' in lowercase. Do not include anything else in the response.\n\n This is an example question: [Question 10: During this hospital stay, did you need help from nurses or other hospital staff in getting to the bathroom or in using a bedpan?a. Yes,b. No (If No, go to question 12)Skipping logic: If the answer is 'No', the survey skips to question 12.] now in this question, if the answer is 'yes' then answer 'no' and if the answer is 'no' then answer '12' as per mentioned in the question
        """
          },
      ]).content[0].text

  return (message)


def extract_number(text):
  # Using \d+ to match one or more digits
  numbers = re.findall(r'\d+', text)
  if numbers:
    # Extracting the first number found and converting it to an integer
    return int(numbers[0])
  else:
    # Returning None if no numbers are found in the text
    return None


# API endpoint to save the current answer and get the next question


@app.route('/text/savegetnextquestion', methods=['POST'])
def save_get_next_question():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')
  responder_id = data.get('responder_id')

  if not responder_id:
    responder_id = '0'
  print(responder_id)
  current_question = data.get('current_question', {})

  if not survey_code or not current_question:
    return jsonify({
        'status':
        'error',
        'error_content':
        'Survey code or current_question not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:

        answer_type = 'text'

        options = current_question.get('options')
        type = current_question.get('type')

        if answer_type == 'voice' and type == 'MCQ':
          answer = answerify(current_question.get('answer'), options)
          print(answer)
        else:
          answer = current_question.get('answer')
        # Save the current answer into the answers table
        cursor.execute(
            """
                    INSERT INTO answers (
                        survey_code, question_number, question, type, answer, responder_id
                    )
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """, (survey_code, current_question.get('question_number'),
                          current_question.get('question'),
                          current_question.get('type'), answer, responder_id))
        connection.commit()

        question = current_question.get('question')
        qn = int(current_question.get('question_number'))
        print("till here")
        cursor.execute(
            """
          SELECT st.skipperlogic
          FROM published AS p
          JOIN skipper_table_questions AS st ON p.outline_id = st.outline_id
          WHERE p.survey_code = %s and st.question_number = %s;
          """, (survey_code, qn))
        #connection.commit()

        skipper_logic = cursor.fetchone()
        #skipper_logic = 'No'
        print("skipper logic is", skipper_logic)

        print("the question is", question)

        print("the answer is", answer)

        skip = whetherskip(question, answer, skipper_logic)

        print("the skip is ", skip)
        skipnumber = extract_number(skip)
        print("the skip number is ", skipnumber)

        # Convert question_number from string to integer, increment it, and get the next question
        next_question_number = str(
            int(current_question.get('question_number')) + 1)

        if skipnumber is not None:
          next_question_number = str(skipnumber)

        print("the next_question_number is ", next_question_number)

        cursor.execute(
            "SELECT outline_id FROM published WHERE survey_code = %s",
            (survey_code, ))
        result = cursor.fetchone()

        if result:
          # Survey code found, now retrieve the first question details from the outline table
          outline_id = result[0]
          # Convert question_number from string to integer, increment it, and get the next question
          # next_question_number = str(int(current_question.get('question_number')) + 1)

          # Check if survey code exists in the published table
          cursor.execute(
              "SELECT question_number, formatted_questions,question_type, options,max_no_of_choices, required,question_text FROM outline WHERE outline_id = %s AND question_number = %s",
              (outline_id, next_question_number))
          next_question_result = cursor.fetchone()

          if next_question_result:
            question_number, formatted_questions, question_type, formatted_options, max_no_of_choices, required, question_text = next_question_result

            if required == 'true':
              skipper = False
            else:
              skipper = True

            # Get the number of rows in the outline table for the given outline_id
            cursor.execute(
                "SELECT COUNT(*) FROM outline WHERE outline_id = %s",
                (outline_id, ))
            number_of_rows_result = cursor.fetchone()
            number_of_rows = number_of_rows_result[
                0] if number_of_rows_result else 0

            # Calculate completion percentage
            completion_percentage = (int(next_question_number) -
                                     1) / number_of_rows
            completion_percentage_str = f"{completion_percentage:.2%}"

            print("answer here is", answer)

            # Return the next question details
            if formatted_questions is not None:
              print("question here is", formatted_questions)
              formatted_questions = dynamic_voice_question(
                  formatted_questions, answer)
              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': formatted_questions,
                  'options': formatted_options,
                  'type': question_type,
                  'completion': completion_percentage_str,
                  'max_no_of_choices': max_no_of_choices,
                  'total_time': '45',
                  'skipquestion': skipper,
                  'responder_id': responder_id
              }), 200
            else:
              print("question here is", question_text)

              question_text = dynamic_voice_question(question_text, answer)

              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': question_text,
                  'options': formatted_options,
                  'type': question_type,
                  'completion': completion_percentage_str,
                  'max_no_of_choices': max_no_of_choices,
                  'total_time': '45',
                  'skipquestion': skipper,
                  'responder_id': responder_id
              }), 200

          # No more questions available
          return jsonify({
              'status': 'success',
              'message': 'No more questions available'
          }), 200

    except Exception as e:
      print(f"Error saving and getting next question: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to save and get next question'
      }), 500

    finally:
      connection.close()

  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


# API endpoint to save the current answer and get the next question
@app.route('/voice/savegetnextquestion', methods=['POST'])
def voice_save_get_next_question():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')
  responder_id = data.get('responder_id')

  if not responder_id:
    responder_id = '0'
  print(responder_id)
  current_question = data.get('current_question', {})

  if not survey_code or not current_question:
    return jsonify({
        'status':
        'error',
        'error_content':
        'Survey code or current_question not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:

        answer_type = 'voice'
        if not answer_type:
          return jsonify({
              'status': 'error',
              'error_content': 'Answer Type  not provided'
          }), 400

        options = current_question.get('options')
        type = current_question.get('type')

        if answer_type == 'voice' and type == 'MCQ':
          answer = answerify(current_question.get('answer'), options)
          print(answer)
        else:
          answer = current_question.get('answer')
        # Save the current answer into the answers table
        cursor.execute(
            """
                    INSERT INTO answers (
                        survey_code, question_number, question, type, answer, responder_id
                    )
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """, (survey_code, current_question.get('question_number'),
                          current_question.get('question'),
                          current_question.get('type'), answer, responder_id))
        connection.commit()

        question = current_question.get('question')
        qn = int(current_question.get('question_number'))
        cursor.execute(
            """
          SELECT st.skipperlogic
          FROM published AS p
          JOIN skipper_table_questions AS st ON p.outline_id = st.outline_id
          WHERE p.survey_code = %s and st.question_number = %s;
          """, (survey_code, qn))

        skipper_logic = cursor.fetchone()

        #connection.commit()
        #skipper_logic = "None"
        print("skipper logic is", skipper_logic)

        print("the question is", question)

        print("the answer is", answer)

        skip = whetherskip(question, answer, skipper_logic)

        print("the skip is ", skip)
        skipnumber = extract_number(skip)
        print("the skip number is ", skipnumber)

        # Convert question_number from string to integer, increment it, and get the next question
        next_question_number = str(
            int(current_question.get('question_number')) + 1)

        if skipnumber is not None:
          next_question_number = str(skipnumber)

        print("the next_question_number is ", next_question_number)

        cursor.execute(
            "SELECT outline_id FROM published WHERE survey_code = %s",
            (survey_code, ))
        result = cursor.fetchone()

        if result:
          # Survey code found, now retrieve the first question details from the outline table
          outline_id = result[0]
          # Convert question_number from string to integer, increment it, and get the next question
          # next_question_number = str(int(current_question.get('question_number')) + 1)

          # Check if survey code exists in the published table
          cursor.execute(
              "SELECT question_number, formatted_questions,question_type, options,max_no_of_choices, required,question_text FROM outline WHERE outline_id = %s AND question_number = %s",
              (outline_id, next_question_number))
          next_question_result = cursor.fetchone()

          if next_question_result:
            question_number, formatted_questions, question_type, formatted_options, max_no_of_choices, required, question_text = next_question_result

            if required == 'true':
              skipper = False
            else:
              skipper = True

            # Get the number of rows in the outline table for the given outline_id
            cursor.execute(
                "SELECT COUNT(*) FROM outline WHERE outline_id = %s",
                (outline_id, ))
            number_of_rows_result = cursor.fetchone()
            number_of_rows = number_of_rows_result[
                0] if number_of_rows_result else 0

            # Calculate completion percentage
            completion_percentage = (int(next_question_number) -
                                     1) / number_of_rows
            completion_percentage_str = f"{completion_percentage:.2%}"

            print("answer here is", answer)

            # Return the next question details
            if formatted_questions is not None:
              print("question here is", formatted_questions)
              if answer_type == 'voice':
                formatted_questions = dynamic_voice_question(
                    formatted_questions, answer)
              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': formatted_questions,
                  'options': formatted_options,
                  'type': question_type,
                  'completion': completion_percentage_str,
                  'max_no_of_choices': max_no_of_choices,
                  'total_time': '45',
                  'skipquestion': skipper,
                  'responder_id': responder_id
              }), 200
            else:
              print("question here is", question_text)
              if answer_type == 'voice':
                question_text = dynamic_voice_question(question_text, answer)

              return jsonify({
                  'status': 'success',
                  'question_number': question_number,
                  'question': question_text,
                  'options': formatted_options,
                  'type': question_type,
                  'completion': completion_percentage_str,
                  'max_no_of_choices': max_no_of_choices,
                  'total_time': '45',
                  'skipquestion': skipper,
                  'responder_id': responder_id
              }), 200

          # No more questions available
          return jsonify({
              'status': 'success',
              'message': 'No more questions available'
          }), 200

    except Exception as e:
      print(f"Error saving and getting next question: {str(e)}")
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to save and get next question'
      }), 500

    finally:
      connection.close()

  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to connect to the database'
    }), 500


@app.route('/getuseroutlines', methods=['GET'])
def get_user_outlines():
  api_secret = request.headers.get('APISECRET')
  jwt_token = request.headers.get('JWTToken')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Get the email from the "signedin" table based on the provided JWT token
        cursor.execute("SELECT email FROM signedin WHERE jwt_token = %s",
                       (jwt_token, ))
        result = cursor.fetchone()

        if not result:
          return jsonify({
              'status': 'error',
              'error_content': 'Invalid JWT token'
          }), 401

        email = result[0]

        # Get a list of all JWT tokens of that particular email
        cursor.execute("SELECT jwt_token FROM signedin WHERE email = %s",
                       (email, ))
        jwt_tokens = [row[0] for row in cursor.fetchall()]

        # Get all the outline_id and status from the "outline" table for each JWT token
        user_outlines = []
        for token in jwt_tokens:
          cursor.execute(
              "SELECT outline_id,status,title,visibility, goal, endafterdate, endafterresponses FROM polls WHERE jwt_token = %s",
              (token, ))
          outlines_data = cursor.fetchall()

          for outline_data in outlines_data:
            input_date = outline_data[5]
            if input_date:
              print("id")
              if isinstance(input_date, (date, datetime)):
                formatted_date = input_date.strftime("%d %b %Y")

            else:
              # If it's a string, then use strptime and strftime
              if input_date:
                try:
                  parsed_date = datetime.strptime(input_date,
                                                  "%a, %d %b %Y %H:%M:%S %Z")
                  formatted_date = parsed_date.strftime("%d %b %Y")

                except ValueError as e:
                  print(f"Error parsing date {input_date}: {e}")

            if outline_data[6]:
              user_outlines.append({
                  'outline_id': outline_data[0],
                  'status': outline_data[1],
                  'title': outline_data[2],
                  'visibility': outline_data[3],
                  'goal': outline_data[4],
                  'endafterdate': formatted_date,
                  'endafterresponses': outline_data[6],
                  'lengthtime': '20 min',
                  'lengthquestions': '7 ques'
              })
            else:
              user_outlines.append({
                  'outline_id': outline_data[0],
                  'status': outline_data[1],
                  'title': outline_data[2],
                  'visibility': 'null',
                  'goal': 'null',
                  'endafterdate': 'null',
                  'endafterresponses': 'null',
                  'lengthtime': 'null',
                  'lengthquestions': 'null'
              })

      return jsonify({
          'status': 'success',
          'user_outlines': user_outlines
      }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to retrieve user outlines. Error: {str(e)}'
      }), 500

    finally:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/getsurveyfromoutline', methods=['POST'])
def getsurveyfromoutline():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  # Validate the access token (implement your token validation logic)
  user = validate_access_token(jwt_token)

  if not user:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid JWT token'
    }), 401

  data = request.get_json()
  outline_id = data.get('idoutline')

  if not outline_id:
    return jsonify({
        'status': 'error',
        'error_content': 'Contents not filled'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Fetch data from the outline table based on outline_id
        cursor.execute(
            """
                    SELECT survey_code FROM published
                    WHERE outline_id = %s
                    """, (outline_id, ))
        survey_code = cursor.fetchall()
        return jsonify({'status': 'success', 'survey_code': survey_code}), 200
    except:

      return jsonify({
          'status':
          'error',
          'error_content':
          'this outline id doesnt have a survey code associated with it as of now'
      }), 500

    finally:
      connection.close()


def add_title(doc, title):
  title_paragraph = doc.add_paragraph()
  title_paragraph.alignment = 1  # Center alignment
  title_run = title_paragraph.add_run(title)
  title_run.bold = True
  title_run.font.size = Pt(18)
  title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black


def add_section(doc, section_title):
  section_heading = doc.add_heading(level=1)
  section_heading.alignment = 0  # Left alignment
  section_heading_run = section_heading.add_run(section_title)
  section_heading_run.bold = True
  section_heading_run.font.size = Pt(14)


def add_instruction(doc, instruction):
  instruction_paragraph = doc.add_paragraph(instruction)
  instruction_paragraph.alignment = 0  # Left alignment


@app.route('/getinsightsfromsurvey', methods=['POST'])
def getinsightsfromsurvey():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  outlineid = data.get('outline_id')

  if not outlineid:
    return jsonify({
        'status': 'error',
        'error_content': 'outline id code not provided'
    }), 400

  connection = get_db_connection()
  if connection:
    try:

      with connection.cursor() as cursor:
        try:
          cursor.execute(
              """
                                    SELECT survey_code FROM published
                                    WHERE outline_id = %s
                                """, (outlineid, ))
          survey = cursor.fetchone()[0]

        except:
          return jsonify({
              'status': 'error',
              'error_content': 'Error fetching survey data'
          }), 500

        # Fetch total number of unique respondents
        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS total_people_reached
                                FROM answers
                                WHERE survey_code = %s
                            """, (survey, ))
        total_people_reached = cursor.fetchone()[0]

        # Fetch number of unique respondents who reached at least the 2nd question
        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS starts
                                FROM answers
                                WHERE survey_code = %s
                                AND CAST(question_number AS INTEGER) >= 2
                            """, (survey, ))
        starts = cursor.fetchone()[0]

        # Fetch number of unique respondents who reached at least the 10th question
        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS completion
                                FROM answers
                                WHERE survey_code = %s
                                AND CAST(question_number AS INTEGER) >= 10
                            """, (survey, ))
        completion = cursor.fetchone()[0]

        # Calculate completion rate in percentage
        completion_rate_percentage = (
            completion /
            total_people_reached) * 100 if total_people_reached > 0 else 0

        # Fetch previous week's data
        last_week_start = datetime.now() - timedelta(days=7)

        # Calculate metrics for the previous week
        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS prev_week_people_reached
                                FROM answers
                                WHERE survey_code = %s
                                AND "timestamp" < %s
                            """, (survey, last_week_start))
        prev_week_people_reached = cursor.fetchone(
        )[0] if cursor.rowcount > 0 else 0

        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS prev_week_starts
                                FROM answers
                                WHERE survey_code = %s
                                AND CAST(question_number AS INTEGER) >= 2
                                AND "timestamp" < %s
                            """, (survey, last_week_start))
        prev_week_starts = cursor.fetchone()[0] if cursor.rowcount > 0 else 0

        cursor.execute(
            """
                                SELECT COUNT(DISTINCT responder_id) AS prev_week_completion
                                FROM answers
                                WHERE survey_code = %s
                                AND CAST(question_number AS INTEGER) >= 10
                                AND "timestamp" < %s
                            """, (survey, last_week_start))
        prev_week_completion = cursor.fetchone(
        )[0] if cursor.rowcount > 0 else 0

        # Calculate increment percentages
        increment_people_reached = (
            (total_people_reached - prev_week_people_reached) /
            prev_week_people_reached
        ) * 100 if prev_week_people_reached > 0 else 0
        increment_starts = (
            (starts - prev_week_starts) /
            prev_week_starts) * 100 if prev_week_starts > 0 else 0
        increment_completion_rate = completion_rate_percentage - (
            (prev_week_completion / prev_week_people_reached) *
            100) if prev_week_people_reached > 0 else 0

        return jsonify({
            'status': 'success',
            'people_reached': total_people_reached,
            'starts': starts,
            'completion_rate': f"{completion_rate_percentage:.2f}%",
            'increment_people_reached': f"{increment_people_reached:.2f}%",
            'increment_starts': f"{increment_starts:.2f}%",
            'increment_completion_rate': f"{increment_completion_rate:.2f}%",
            'average_time_to_complete': '00:00',
            'ai_time_to_complete': '13:00',
        }), 200

    except Exception as e:
      print("An error occurred:", e)
      return jsonify({
          'status':
          'error',
          'error_content':
          'An error occurred while processing the request'
      }), 500
    finally:
      connection.close()
  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to establish a database connection'
    }), 500


@app.route('/export-poll', methods=['POST'])
def export_poll():
  data = request.json
  if not data:
    return 'No data provided', 400

  # Create a Word document
  doc = Document()

  # Add Title
  title = 'Sample Poll'
  if title:
    add_title(doc, title)

  # Add Goal
  goal = 'To gather feedback from users'

  if goal:
    add_section(doc, 'Goal:')
    add_instruction(doc, goal)

  # Add Instruction
  instruction = 'Please answer the following questions'

  if instruction:
    add_section(doc, 'Instruction:')
    add_instruction(doc, instruction)

  # Add Introduction
  introduction = 'Thank you for participating in this poll'

  if introduction:
    add_section(doc, 'Introduction:')
    add_instruction(doc, introduction)

  # Add Questions
  questions = [{
      "text": "What is your favorite color?"
  }, {
      "text": "What is your favorite food?",
      "options": ["Pizza", "Sushi", "Burgers"]
  }, {
      "text": "What is your age?"
  }]
  #questions = data.get('questions')
  if questions:
    add_section(doc, 'Questions:')
    for idx, question in enumerate(questions, start=1):
      question_text = question['text']
      options = question.get('options')
      question_paragraph = doc.add_paragraph(f'{idx}. {question_text}')
      if options:
        question_paragraph.style = 'ListBullet'
        for opt_idx, option in enumerate(options, start=1):
          question_paragraph.add_run(f'\n   {opt_idx}. {option}')

  # Save the document to a BytesIO buffer
  temp_buffer = io.BytesIO()
  doc.save(temp_buffer)
  temp_buffer.seek(0)

  # Return the document as a file attachment
  return send_file(
      temp_buffer,
      as_attachment=True,
      mimetype=
      'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
      download_name='poll_document.docx')


def duplicate_outline_data(outline_id):
  # Connect to your database
  connection = get_db_connection()

  cursor = connection.cursor()

  # Get existing data
  cursor.execute("SELECT * FROM polls WHERE outline_id = %s", (outline_id, ))
  existing_poll = cursor.fetchone()

  # If no existing data found, return error
  if not existing_poll:
    return jsonify({"error": "Outline ID not found"}), 404

  # Generate new outline ID
  new_outline_id = generate_outline_id(cursor)

  newtitle = "copy of " + existing_poll[2]

  # Duplicate data in polls table
  cursor.execute(
      "INSERT INTO polls (outline_id, title, goal, endafterdate, endafterresponses, email, geography, education, industry, visibility, status, jwt_token, introduction, instruction, workspace) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
      (new_outline_id, newtitle, existing_poll[3], existing_poll[5],
       existing_poll[6], existing_poll[7], existing_poll[8], existing_poll[9],
       existing_poll[10], existing_poll[11], "Draft Ready", existing_poll[13],
       existing_poll[14], existing_poll[15], existing_poll[16]))

  # Duplicate data in outline table
  cursor.execute("SELECT * FROM outline WHERE outline_id = %s", (outline_id, ))
  existing_outline_data = cursor.fetchall()
  for row in existing_outline_data:
    cursor.execute(
        "INSERT INTO outline (title, question_number, importance, required, instruction, dynamic_followup, objective, max_no_of_questions, keywords_to_probe, things_to_avoid, example_questions, question_text, question_type, branching, outline_id, status, allow_others, max_no_of_choices, question_full, formatted_options, introduction) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
        (row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
         row[9], row[10], row[11], row[12], row[13], row[14], new_outline_id,
         "Draft Ready", row[18], row[19], row[20], row[22], row[21]))

  # Commit changes and close connection
  connection.commit()
  cursor.close()
  connection.close()

  return new_outline_id


@app.route('/duplicateoutline', methods=['POST'])
def duplicate_outline():
  data = request.json
  if 'idoutline' not in data:
    return jsonify({"error": "Outline ID not provided"}), 400

  outline_id = data['idoutline']
  new_outline_id = duplicate_outline_data(outline_id)

  return jsonify({"new_outline_id": new_outline_id}), 200


@app.route('/deleteoutline', methods=['POST'])
def delete_outline():
  data = request.json
  if 'idoutline' not in data:
    return jsonify({"error": "Outline ID not provided"}), 400

  # the table name is polls
  # i want to delete all from polls where outline_id = idoutline
  outline_id = data['idoutline']

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      # Define the SQL statement to delete rows with the given outline_id
      delete_query = "DELETE FROM polls WHERE outline_id = %s"

      # Execute the SQL statement with the provided outline_id
      cursor.execute(delete_query, (outline_id, ))

      connection.commit()

      cursor.close()
      connection.close()

      return jsonify({'status': 'success'}), 200
    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({'status': 'error', "error": str(e)}), 500
    finally:
      cursor.close()
      connection.close()


@app.route('/insert_workspace', methods=['POST'])
def insert_workspace():
  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Get the email from the "signedin" table based on the provided JWT token
        cursor.execute("SELECT email FROM signedin WHERE jwt_token = %s",
                       (jwt_token, ))
        result = cursor.fetchone()

        if not result:
          return jsonify({
              'status': 'error',
              'error_content': 'Invalid JWT token'
          }), 401

        email = result[0]

        data = request.get_json()
        workspace = data.get('workspace')

        if not workspace:
          return jsonify({
              'status':
              'error',
              'error_content':
              'Workspace is required in the request body'
          }), 400

        # Insert email and workspace into the workspace table
        cursor.execute(
            """
                    INSERT INTO workspace (email, workspace)
                    VALUES (%s, %s)
                    """, (email, workspace))
        connection.commit()

        return jsonify({
            'status': 'success',
            'message': 'Workspace added successfully'
        }), 201

    except psycopg2.Error as e:
      print("Error executing SQL query:", e)
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to insert workspace'
      }), 500

    finally:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Failed to insert workspace'
  }), 500


@app.route('/get_workspaces', methods=['GET'])
def get_workspaces():
  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Get the email from the "signedin" table based on the provided JWT token
        cursor.execute("SELECT email FROM signedin WHERE jwt_token = %s",
                       (jwt_token, ))
        result = cursor.fetchone()

        if not result:
          return jsonify({
              'status': 'error',
              'error_content': 'Invalid JWT token'
          }), 401

        email = result[0]

        # Fetch all workspaces associated with the email
        cursor.execute(
            """
                    SELECT workspace FROM workspace
                    WHERE email = %s
                    """, (email, ))
        workspaces = cursor.fetchall()

        workspaces_data = []

        cursor.execute("SELECT jwt_token FROM signedin WHERE email = %s",
                       (email, ))
        tokens = cursor.fetchall()

        for workspace in workspaces:
          workspace_name = workspace[0]

          polls_count = 0
          for token in tokens:
            # Count the number of polls associated with the workspace
            cursor.execute(
                """
                          SELECT COUNT(*) FROM polls
                          WHERE workspace = %s AND jwt_token = %s
                          """, (
                    workspace_name,
                    token,
                ))

            polls_count = polls_count + cursor.fetchone()[0]

          workspaces_data.append({
              'workspace': workspace_name,
              'polls_count': polls_count
          })

        # Check if "default workspace" is already in the list
        default_workspace_count = 0
        for workspace_data in workspaces_data:
          if workspace_data['workspace'] == 'My workspace':
            default_workspace_count = workspace_data['polls_count']
            break

        if default_workspace_count == 0:
          default_workspace_count = 0
          for token in tokens:
            # Count the number of polls associated with the default workspace
            cursor.execute(
                """
                          SELECT COUNT(*) FROM polls
      WHERE workspace = 'My workspace' AND jwt_token = %s
                          """, (token, ))
            default_workspace_count = default_workspace_count + cursor.fetchone(
            )[0]

          workspaces_data.append({
              'workspace': 'My workspace',
              'polls_count': default_workspace_count
          })

        return jsonify({
            'status': 'success',
            'workspaces': workspaces_data
        }), 200

    except psycopg2.Error as e:
      print("Error executing SQL query:", e)
      return jsonify({
          'status': 'error',
          'error_content': 'Failed to fetch workspaces'
      }), 500

    finally:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Failed to fetch workspaces'
  }), 500


@app.route('/getworkspaceoutlines', methods=['POST'])
def get_workspace_outlines():
  api_secret = request.headers.get('APISECRET')
  jwt_token = request.headers.get('JWTToken')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  workspace = data.get('workspace')

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Get the email from the "signedin" table based on the provided JWT token
        cursor.execute("SELECT email FROM signedin WHERE jwt_token = %s",
                       (jwt_token, ))
        result = cursor.fetchone()

        if not result:
          return jsonify({
              'status': 'error',
              'error_content': 'Invalid JWT token'
          }), 401

        email = result[0]

        # Get a list of all JWT tokens of that particular email
        cursor.execute("SELECT jwt_token FROM signedin WHERE email = %s",
                       (email, ))
        jwt_tokens = [row[0] for row in cursor.fetchall()]

        # Get all the outline_id and status from the "outline" table for each JWT token
        user_outlines = []
        for token in jwt_tokens:
          cursor.execute(
              "SELECT outline_id,status,title,visibility, goal, endafterdate, endafterresponses FROM polls WHERE jwt_token = %s AND workspace = %s",
              (
                  token,
                  workspace,
              ))
          outlines_data = cursor.fetchall()

          print(outlines_data)

          for outline_data in outlines_data:
            input_date = outline_data[5]
            if input_date:
              print("id")
              if isinstance(input_date, (date, datetime)):
                formatted_date = input_date.strftime("%d %b %Y")

            else:
              # If it's a string, then use strptime and strftime
              if input_date:
                try:
                  parsed_date = datetime.strptime(input_date,
                                                  "%a, %d %b %Y %H:%M:%S %Z")
                  formatted_date = parsed_date.strftime("%d %b %Y")

                except ValueError as e:
                  print(f"Error parsing date {input_date}: {e}")

            if outline_data[6]:
              if outline_data[1] == "Processing":
                user_outlines.append({
                    'outline_id': outline_data[0],
                    'status': outline_data[1],
                    'title': outline_data[2],
                    'visibility': outline_data[3],
                    'goal': outline_data[4],
                    'endafterdate': formatted_date,
                    'endafterresponses': outline_data[6],
                    'lengthtime': '0 min',
                    'lengthquestions': '0 ques',
                    'completion': '0%'
                })
              else:
                id = outline_data[0]
                cursor.execute(
                    "SELECT max_no_of_questions from outline WHERE outline_id = %s",
                    (id, ))
                max_data = cursor.fetchone()
                max = int(max_data[0])
                #rtime = random.randint(20, 40)

                rleng = max
                rtime = max * 2

                user_outlines.append({
                    'outline_id': outline_data[0],
                    'status': outline_data[1],
                    'title': outline_data[2],
                    'visibility': outline_data[3],
                    'goal': outline_data[4],
                    'endafterdate': formatted_date,
                    'endafterresponses': outline_data[6],
                    'lengthtime': f'{rtime} min',
                    'lengthquestions': f'{rleng} ques',
                    'completion': '0%'
                })
            else:
              user_outlines.append({
                  'outline_id': outline_data[0],
                  'status': outline_data[1],
                  'title': outline_data[2],
                  'visibility': 'null',
                  'goal': 'null',
                  'endafterdate': 'null',
                  'endafterresponses': 'null',
                  'lengthtime': 'null',
                  'lengthquestions': 'null',
                  'completion': '0%'
              })

      return jsonify({
          'status': 'success',
          'user_outlines': user_outlines
      }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to retrieve user outlines. Error: {str(e)}'
      }), 500

    finally:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/deleteworkspace', methods=['POST'])
def deleteworkspace():
  api_secret = request.headers.get('APISECRET')
  jwt_token = request.headers.get('JWTToken')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  workspace = data.get('workspace')

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        # Get the email from the "signedin" table based on the provided JWT token
        cursor.execute("SELECT email FROM signedin WHERE jwt_token = %s",
                       (jwt_token, ))
        result = cursor.fetchone()

        if not result:
          return jsonify({
              'status': 'error',
              'error_content': 'Invalid JWT token'
          }), 401

        email = result[0]

        # Get a list of all JWT tokens of that particular email
        cursor.execute(
            "select workspace FROM workspace WHERE email = %s and workspace = %s",
            (
                email,
                workspace,
            ))

        wsresult = cursor.fetchone()

        if not wsresult:
          return jsonify({
              'status': 'error',
              'error_content': 'Workspace or email doesnt exist'
          }), 401
        workspace = wsresult[0]
        print(email)
        print(workspace)

        cursor.execute(
            "Delete FROM workspace WHERE email = %s and workspace = %s", (
                email,
                workspace,
            ))
        connection.commit()
      return jsonify({
          'status': 'success',
          'content': 'Workspace deleted successfully'
      }), 200

    except Exception as e:
      print(f"Error: {str(e)}")
      return jsonify({
          'status':
          'error',
          'error_content':
          f'Failed to delete workspace. Error: {str(e)}'
      }), 500

    finally:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


@app.route('/getresponsefromsurvey', methods=['POST'])
def getresponsefromsurvey():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  outlineid = data.get('outline_id')

  if not outlineid:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        cursor.execute(
            """
                    SELECT survey_code FROM published
                    WHERE outline_id = %s
                """, (outlineid, ))
        survey = cursor.fetchone()[0]
        print(survey)
        # Fetch responses for the given survey code
        cursor.execute(
            """
                    SELECT question_number, question,type, answer, COUNT(*) AS count
                    FROM answers
                    WHERE survey_code = %s
                    GROUP BY question_number, question,type, answer
                    ORDER BY question_number, answer
                """, (survey, ))
        rows = cursor.fetchall()

        response = {}

        for row in rows:
          question_number = row[0]
          question = row[1]
          type = row[2]
          answer = row[3]
          count = row[4]

          # Accumulate count for total count

          # If question_number is not already in the response dictionary, initialize it
          if question_number not in response:
            response[question_number] = {
                'question': question,
                'type': type,
                'answers': []
            }

          # Add answer count to the response
          response[question_number]['answers'].append({
              'answer': answer,
              'count': count
          })
        #response['total_count'] = total_count
        # Calculate percentages for each answer count
        for question_number in response:
          total_count = 0  # Initialize total count
          total_responses = sum(
              answer['count']
              for answer in response[question_number]['answers'])

          for answer_info in response[question_number]['answers']:
            answer_info['percentage'] = (answer_info['count'] /
                                         total_responses) * 100
          total_count += total_responses

          print(total_count)
          response[question_number]['total_count'] = total_count
          response[question_number]['user_custom'] = 'user'

        return jsonify({'status': 'success', 'data': response}), 200

    except Exception as e:
      print("An error occurred:", e)
      return jsonify({
          'status':
          'error',
          'error_content':
          'An error occurred while processing the request'
      }), 500
    finally:
      connection.close()
  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to establish a database connection'
    }), 500


@app.route('/get_drop_off_amounts', methods=['POST'])
def get_drop_off_amounts():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  jwt_token = request.headers.get('JWTToken')

  if not jwt_token:
    return jsonify({
        'status': 'error',
        'error_content': 'JWT token is required'
    }), 401

  data = request.get_json()
  outlineid = data.get('outline_id')

  if not outlineid:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code not provided'
    }), 400

  connection = get_db_connection()

  if connection:
    try:
      with connection.cursor() as cursor:
        try:
          cursor.execute(
              """
                      SELECT survey_code FROM published
                      WHERE outline_id = %s
                  """, (outlineid, ))
          survey = cursor.fetchone()[0]
          print(survey)
        except:
          drop_off_amounts = {
              f'Q{i}': 0
              for i in range(1, 25)
          }  # Assuming 10 questions
          highest_question_number = 25

          for q in range(1, highest_question_number + 1):
            drop_off_amounts[f'Q{q}'] = 0

          return jsonify({
              'status': 'success',
              'drop_off_amounts': drop_off_amounts
          }), 200
        # Fetch the highest question number answered by each unique responder
        cursor.execute(
            """
                    SELECT CAST(responder_id AS INTEGER) AS responder_id,
                           CAST(MAX(question_number) AS INTEGER) AS highest_question_number
                    FROM (
                        SELECT responder_id, question_number
                        FROM answers
                        WHERE survey_code = %s
                        ORDER BY responder_id, CAST(question_number AS INTEGER) DESC
                    ) AS ordered_answers
                    GROUP BY responder_id
                """, (survey, ))
        # Fetch the rows from the database
        rows = cursor.fetchall()

        # Convert string values to integers
        rows = [(int(responder_id), int(highest_question_number))
                for responder_id, highest_question_number in rows]

        print(rows)

        # Initialize a dictionary to store drop-off amounts for each question
        drop_off_amounts = {
            f'Q{i}': 0
            for i in range(1, 25)
        }  # Assuming 10 questions

        # Calculate drop-off amounts
        for row in rows:
          try:
            highest_question_number = int(row[1])
          except ValueError:
            print(f"Invalid question number: {row[1]}")
            continue

          for q in range(1, highest_question_number + 1):
            drop_off_amounts[f'Q{q}'] += 1

        return jsonify({
            'status': 'success',
            'drop_off_amounts': drop_off_amounts
        }), 200

    except Exception as e:
      print("An error occurred:", e)
      return jsonify({
          'status':
          'error',
          'error_content':
          'An error occurred while processing the request'
      }), 500
    finally:
      connection.close()
  else:
    return jsonify({
        'status': 'error',
        'error_content': 'Unable to establish a database connection'
    }), 500


import logging


@app.route('/savefeedback', methods=['POST'])
def save_feedback():
  api_secret = request.headers.get('APISECRET')

  if api_secret != APISECRET:
    return jsonify({
        'status': 'error',
        'error_content': 'Invalid API secret'
    }), 401

  data = request.get_json()
  survey_code = data.get('survey_code')
  responder_id = data.get('responder_id', '0')
  feedback_text = data.get('feedback_text')
  feedback_additional = data.get("feedback_additional")

  if not survey_code or not feedback_text:
    return jsonify({
        'status': 'error',
        'error_content': 'Survey code or feedback not provided'
    }), 400

  try:
    connection = get_db_connection()
    if connection:
      with connection.cursor() as cursor:
        cursor.execute(
            """
                    INSERT INTO feedback (survey_code, feedback_text, feedback_additional, responder_id)
                    VALUES (%s, %s, %s, %s)
                    """,
            (survey_code, feedback_text, feedback_additional, responder_id))
      connection.commit()

      return jsonify({'status': 'success', 'message': 'Feedback saved'}), 200

  except Exception as e:
    logging.error(f"Error saving feedback: {str(e)}")
    return jsonify({
        'status': 'error',
        'error_content': 'Failed to save feedback'
    }), 500

  finally:
    if connection:
      connection.close()

  return jsonify({
      'status': 'error',
      'error_content': 'Unable to connect to the database'
  }), 500


def resultify(survey_code, section_title, requirement, question_answer_pairs):
  response = openai.ChatCompletion.create(
      model="gpt-4-1106-preview",
      messages=[{
          "role":
          "system",
          "content":
          f"Generate a detailed summary for the section titled '{section_title}' of the survey with code '{survey_code}' based on the requirement and questions with answers provided below. The summary should include:\n\n"
          f"1. An overall summary of the feedback for the section.\n"
          f"2. A list of strengths mentioned in the responses.\n"
          f"3. Areas for development highlighted by the respondents.\n"
          f"4. A visual representation of feedback trends in tabular form, showing the number of mentions for strengths and areas for development.\n\n"
          f"5. A word cloud data of the insights which is basically a list of words and their sizes"
          f"Requirement: {requirement}\n\n"
          f"Questions and Answers:\n{question_answer_pairs}"
      }])
  results = response['choices'][0]['message']['content']
  return results


@app.route('/survey', methods=['POST'])
def get_survey_responses():
  data = request.get_json()
  survey_code = data.get('survey_code')
  section_title = data.get('section_title')

  if not survey_code or not section_title:
    return jsonify({'error':
                    'survey_code and section_title are required'}), 400

  requirements = {
      'Vision and Direction':
      """
            Primarily if his audience understands the vision and strategic direction. I want to know if they
            are clear how their role/job plays a part in executing this vision. During the interviews I asked
            each participant what the vision and strategy was and although their first response was that he
            was clear, no one could answer the question. If he does not, the leader could not articulate the
            vision and strategy and/or they could not articulate the purpose of their role to execute this
            vision. Secondarily I’m looking for his clarity and how he communicates. Is he shifting his style to
            different audiences?
        """,
      'Leadership Style':
      """
            I am looking to see if he is a visionary leader who can effectively communicate to all audiences
            and adjust accordingly. Does he act decisively and does he have self and situational awareness
            to understand the needs, inefficiencies or concerns of the employees or staff. Does he take
            information from staff or non-physicians and integrate ideas from everyone. Does he foster a
            growth mindset illustrated by his recognition of trying and learning versus solely on
            performance excellence and perfection. Does he balance goals with strategic thinking, to
            ensure he’s staying on track with the current goals and not overwhelming his leaders with new
            ideas frequently. I am looking for the impact of his style on the organization such as, inefficient
            work, lack of clarity, no sense of safety in speaking up in meetings, in effective use of chain of
            command, inequity with reporting structures. Some of the VP’s report to the CEO and some
            report to the COO. This provides an inequitable level of knowledge and understanding of the
            strategies to different Executive Team members. The reporting structure reflects the CEO’s interests.
        """,
      'Decision-Making':
      """
            Does he respond to a crisis clearly and immediately? Does he look at decisions from all aspects,
            not just analytical (or what the CEO style is) but also emotional, systems or structural and does
            he apply a nuance to his decision making that takes in more than just data. Does he have a calm
            about him in times of uncertainty and challenge and applies analytically thinking, risk assessments,
            alignment with his Executive Team or Board Stakeholders and seeks and applies the feedback he learns.
            Does he know how to prioritize, understand the pace of change and when to pull back vs.
            continuing to add projects to the team’s workload.
        """,
      'Communication and Transparency':
      """
            I’m looking for: 1. Visual Management on performance and progress 2. Regular, consistent and impactful communication
            to stakeholder audiences such as Board Members, Executive Team, and Staff. 3. Does he apply communication externally leveraging
            town hall formats, social media, email, and other media to ensure he’s clear about the directions and decisions. 4. Is he inclusive, and
            aware of cultural differences? 5. Does he directly deal with conflict and handle issues internal in the organization that are creating
            confusion, and inefficiencies? 6. Is he communicating publicly to external and internal audiences comfortably and often? 7. Is he
            compelling and inspiring? 8. Is he aware of his non-verbal communication? 9. Are decisions being made in a group setting or off line
            in 1x1 meetings without clarity or communication to the large leadership team?
        """
  }

  requirement = requirements.get(section_title)
  if not requirement:
    return jsonify({'error': 'Invalid section title'}), 400

  connection = get_db_connection()
  if connection:
    try:
      cursor = connection.cursor()

      query = """
                SELECT question_text, answer_text 
                FROM survey_responses 
                WHERE survey_code = %s AND section_title = %s
            """
      cursor.execute(query, (survey_code, section_title))
      responses = cursor.fetchall()

      question_answer_pairs = [
          f"{question} {answer}" for question, answer in responses
      ]

      # Print the question-answer pairs to the console
      for pair in question_answer_pairs:
        print(pair)

      cursor.close()
      connection.close()

      result = resultify(survey_code, section_title, requirement,
                         question_answer_pairs)

      print(result)
      return jsonify({'message': 'success', 'result': result}), 200
    except Exception as e:
      print(f"Error: Unable to fetch data. {str(e)}")
      return jsonify({'error': 'Unable to fetch data'}), 500
  else:
    return jsonify({'error': 'Unable to connect to the database'}), 500


if __name__ == '__main__':
  app.run(host='0.0.0.0', port=5000)
